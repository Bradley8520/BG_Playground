from docx import Document
from docx.shared import Pt
import re

src = r'c:\Users\bg763c\OneDrive - AT&T Services, Inc\Documents\VS Code\Git Test\Taxonomy_Skill_Master_Reference_DRAFT.md'
out = r'c:\Users\bg763c\OneDrive - AT&T Services, Inc\Documents\VS Code\Git Test\Taxonomy_Skill_Master_Reference_DRAFT.docx'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

with open(src, encoding='utf-8') as f:
    lines = f.readlines()

in_code = False
code_lines = []

def add_code_block(document, block_lines):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('\n'.join(block_lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def add_inline(p, text):
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            p.add_run(part)

fence = '```'

for line in lines:
    line = line.rstrip('\n')

    if line.startswith(fence):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            add_code_block(doc, code_lines)
            in_code = False
            code_lines = []
        continue

    if in_code:
        code_lines.append(line)
        continue

    if line.startswith('# ') and not line.startswith('## '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        add_inline(p, line[2:])
    elif re.match(r'^\d+\. ', line):
        p = doc.add_paragraph(style='List Number')
        add_inline(p, re.sub(r'^\d+\. ', '', line))
    elif line.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(line[2:])
        run.italic = True
    elif line.startswith('---'):
        doc.add_paragraph('─' * 60)
    elif line.startswith('|'):
        p = doc.add_paragraph(line)
        if p.runs:
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
    elif line.strip() == '':
        doc.add_paragraph('')
    else:
        p = doc.add_paragraph()
        add_inline(p, line)

doc.save(out)
print('Saved:', out)
